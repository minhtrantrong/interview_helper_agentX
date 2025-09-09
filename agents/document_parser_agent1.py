"""
Industrial-Grade Document Parser
Handles PDFs, Word docs, text files, images, and more with robust error handling
"""
import os
import re
import logging
from typing import Optional, Dict, List, Tuple
from dataclasses import dataclass
import PyPDF2
import pdfplumber
import docx
import chardet
from striprtf.striprtf import rtf_to_text
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)

@dataclass
class ParseResult:
    success: bool
    text: str
    metadata: Dict
    method_used: str
    error: Optional[str] = None

class DocumentParser:
    """Robust document parser that handles multiple formats and edge cases"""
    
    def __init__(self):
        # Check if tesseract is available for OCR
        self.ocr_available = self._check_tesseract()
        
        # Supported file extensions
        self.supported_extensions = {
            '.pdf': self.parse_pdf,
            '.doc': self.parse_word,
            '.docx': self.parse_word,
            '.txt': self.parse_text,
            '.rtf': self.parse_rtf,
            '.png': self.parse_image,
            '.jpg': self.parse_image,
            '.jpeg': self.parse_image,
            '.tiff': self.parse_image,
            '.bmp': self.parse_image
        }
        
        logger.info(f"📄 Document Parser initialized. OCR: {'✅' if self.ocr_available else '❌'}")
    
    def _check_tesseract(self) -> bool:
        """Check if tesseract OCR is available"""
        try:
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            logger.warning("❌ Tesseract OCR not available. Install with: brew install tesseract")
            return False
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove weird Unicode characters
        text = text.encode('utf-8', errors='ignore').decode('utf-8')
        
        # Remove control characters but keep newlines and tabs
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        # Normalize line breaks
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _detect_encoding(self, file_path: str) -> str:
        """Detect file encoding"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8') or 'utf-8'
        except Exception:
            return 'utf-8'
    
    def parse_pdf(self, file_path: str) -> ParseResult:
        """Parse PDF with multiple strategies"""
        methods_tried = []
        
        # Strategy 1: pdfplumber (best for formatted PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                tables_found = 0
                
                for page in pdf.pages:
                    # Extract text
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        tables_found += len(tables)
                        for table in tables:
                            # Convert table to text
                            table_text = "\n".join(["\t".join([cell or "" for cell in row]) for row in table])
                            text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
                
                full_text = "\n".join(text_parts)
                cleaned_text = self._clean_text(full_text)
                
                if cleaned_text and len(cleaned_text.strip()) > 50:
                    return ParseResult(
                        success=True,
                        text=cleaned_text,
                        metadata={
                            "pages": len(pdf.pages),
                            "tables_found": tables_found,
                            "file_size": os.path.getsize(file_path),
                            "has_tables": tables_found > 0
                        },
                        method_used="pdfplumber"
                    )
                methods_tried.append("pdfplumber")
                
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
            methods_tried.append(f"pdfplumber (failed: {str(e)[:50]})")
        
        # Strategy 2: PyPDF2 (fallback for encrypted or simple PDFs)
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                full_text = "\n".join(text_parts)
                cleaned_text = self._clean_text(full_text)
                
                if cleaned_text and len(cleaned_text.strip()) > 50:
                    return ParseResult(
                        success=True,
                        text=cleaned_text,
                        metadata={
                            "pages": len(pdf_reader.pages),
                            "file_size": os.path.getsize(file_path),
                            "encrypted": pdf_reader.is_encrypted
                        },
                        method_used="PyPDF2"
                    )
                methods_tried.append("PyPDF2")
                
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")
            methods_tried.append(f"PyPDF2 (failed: {str(e)[:50]})")
        
        # Strategy 3: OCR if available (for image-based PDFs)
        if self.ocr_available:
            try:
                # Convert PDF pages to images and OCR them
                import fitz  # PyMuPDF - would need to add this if we want it
                logger.warning("OCR strategy not fully implemented - requires PyMuPDF")
                methods_tried.append("OCR (not implemented)")
            except ImportError:
                methods_tried.append("OCR (PyMuPDF not available)")
        
        # All methods failed
        return ParseResult(
            success=False,
            text="",
            metadata={"methods_tried": methods_tried, "file_size": os.path.getsize(file_path)},
            method_used="none",
            error="All PDF parsing methods failed"
        )
    
    def parse_word(self, file_path: str) -> ParseResult:
        """Parse Word documents (.doc/.docx)"""
        try:
            # Only .docx is supported by python-docx
            if file_path.lower().endswith('.docx'):
                doc = docx.Document(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                tables_found = 0
                for table in doc.tables:
                    tables_found += 1
                    table_text = []
                    for row in table.rows:
                        row_text = "\t".join([cell.text.strip() for cell in row.cells])
                        table_text.append(row_text)
                    
                    if table_text:
                        text_parts.append(f"\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")
                
                full_text = "\n".join(text_parts)
                cleaned_text = self._clean_text(full_text)
                
                return ParseResult(
                    success=True,
                    text=cleaned_text,
                    metadata={
                        "paragraphs": len(doc.paragraphs),
                        "tables_found": tables_found,
                        "file_size": os.path.getsize(file_path),
                        "has_tables": tables_found > 0
                    },
                    method_used="python-docx"
                )
            else:
                # .doc files are not supported
                return ParseResult(
                    success=False,
                    text="",
                    metadata={"file_size": os.path.getsize(file_path)},
                    method_used="none",
                    error=".doc files not supported, please convert to .docx"
                )
                
        except Exception as e:
            return ParseResult(
                success=False,
                text="",
                metadata={"file_size": os.path.getsize(file_path)},
                method_used="python-docx",
                error=f"Word document parsing failed: {str(e)}"
            )
    
    def parse_text(self, file_path: str) -> ParseResult:
        """Parse plain text files with encoding detection"""
        try:
            # Detect encoding
            encoding = self._detect_encoding(file_path)
            
            # Try detected encoding first
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
                encoding = "utf-8 (with errors replaced)"
            
            cleaned_text = self._clean_text(text)
            
            return ParseResult(
                success=True,
                text=cleaned_text,
                metadata={
                    "encoding_detected": encoding,
                    "file_size": os.path.getsize(file_path),
                    "line_count": text.count('\n') + 1
                },
                method_used="text_reader"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                text="",
                metadata={"file_size": os.path.getsize(file_path)},
                method_used="text_reader",
                error=f"Text file parsing failed: {str(e)}"
            )
    
    def parse_rtf(self, file_path: str) -> ParseResult:
        """Parse RTF (Rich Text Format) files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                rtf_content = f.read()
            
            # Convert RTF to plain text
            text = rtf_to_text(rtf_content)
            cleaned_text = self._clean_text(text)
            
            return ParseResult(
                success=True,
                text=cleaned_text,
                metadata={
                    "file_size": os.path.getsize(file_path),
                    "original_rtf_size": len(rtf_content)
                },
                method_used="striprtf"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                text="",
                metadata={"file_size": os.path.getsize(file_path)},
                method_used="striprtf",
                error=f"RTF parsing failed: {str(e)}"
            )
    
    def parse_image(self, file_path: str) -> ParseResult:
        """Parse images using OCR"""
        if not self.ocr_available:
            return ParseResult(
                success=False,
                text="",
                metadata={"file_size": os.path.getsize(file_path)},
                method_used="none",
                error="OCR not available. Install tesseract: brew install tesseract"
            )
        
        try:
            # Open and process image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use OCR to extract text
            text = pytesseract.image_to_string(image, lang='eng')
            cleaned_text = self._clean_text(text)
            
            return ParseResult(
                success=True,
                text=cleaned_text,
                metadata={
                    "image_size": image.size,
                    "image_mode": image.mode,
                    "file_size": os.path.getsize(file_path)
                },
                method_used="pytesseract_ocr"
            )
            
        except Exception as e:
            return ParseResult(
                success=False,
                text="",
                metadata={"file_size": os.path.getsize(file_path)},
                method_used="pytesseract_ocr",
                error=f"OCR failed: {str(e)}"
            )
    
    def parse_document(self, file_path: str) -> ParseResult:
        """Main parsing method that routes to appropriate parser"""
        if not os.path.exists(file_path):
            return ParseResult(
                success=False,
                text="",
                metadata={},
                method_used="none",
                error=f"File not found: {file_path}"
            )
        
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        if ext not in self.supported_extensions:
            return ParseResult(
                success=False,
                text="",
                metadata={"file_extension": ext},
                method_used="none",
                error=f"Unsupported file type: {ext}. Supported: {list(self.supported_extensions.keys())}"
            )
        
        # Route to appropriate parser
        parser_method = self.supported_extensions[ext]
        
        try:
            logger.info(f"📄 Parsing {file_path} with {parser_method.__name__}")
            result = parser_method(file_path)
            
            if result.success:
                logger.info(f"✅ Successfully parsed {file_path} using {result.method_used}")
                logger.info(f"   📊 Extracted {len(result.text)} characters")
            else:
                logger.warning(f"❌ Failed to parse {file_path}: {result.error}")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Unexpected error parsing {file_path}: {e}", exc_info=True)
            return ParseResult(
                success=False,
                text="",
                metadata={"file_extension": ext},
                method_used=parser_method.__name__,
                error=f"Unexpected parsing error: {str(e)}"
            )
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_extensions.keys())
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate if a file can be processed"""
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        if os.path.getsize(file_path) == 0:
            return False, "File is empty"
        
        if os.path.getsize(file_path) > 50 * 1024 * 1024:  # 50MB limit
            return False, "File too large (>50MB)"
        
        _, ext = os.path.splitext(file_path.lower())
        if ext not in self.supported_extensions:
            return False, f"Unsupported format: {ext}"
        
        return True, "File is valid"